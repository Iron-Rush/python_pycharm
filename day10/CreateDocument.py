# CreateDocument.py
from docx import Document
document = Document()
document.add_heading('Document Title', 0)
p = document.add_paragraph('A plain paragtaph having some')
document.add_page_break()
document.save('demo.docx')